import sys
import os
import re
import glob
from datetime import datetime
import logging
import logging.config
import pytz
import shutil
from docx import Document
from docx.shared import Inches
import numpy as np
import json
import matplotlib.pyplot as plt
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

process_start_time = datetime.now(pytz.timezone("Asia/Shanghai"))
result_folder = './logs/' + process_start_time.strftime("%Y%m%d_%H%M%S") + '{desc}'
result_folder_ini = './logs/' + process_start_time.strftime("%Y%m%d_%H%M%S") + '{desc}'

class JSONFormatter(logging.Formatter):
    def format(self, record):
        log_record = {
            'time': self.formatTime(record, self.datefmt),
            'message': record.msg,
        }
        return json.dumps(log_record, indent=5)


class ReportCreator:
    def __init__(self, paras, title, out_path, method_type='p') -> None:
        self.title = title
        self.method_type = method_type
        self.exp_output_path = out_path
        self.paras = paras

        # identify the pop_size and pop_number
        self.pop_number = None
        self.pop_size = None
        self.algs_number = None
        self.ini_pop_size = None

        if method_type == 'p':
            all_algorithm_files = glob.glob(os.path.join(self.exp_output_path + '/history_result', '*.json'))
            all_algorithm_files = sorted(all_algorithm_files, key=lambda x: int(re.search(r'\d+', os.path.basename(x)).group()))
            # initial population may not be full
            with open(all_algorithm_files[0]) as file:
                count_algs = json.load(file)['message']
                self.ini_pop_size = len(count_algs)
            with open(all_algorithm_files[-1]) as file:
                count_algs = json.load(file)['message']
                self.pop_size = len(count_algs)
            self.pop_number = len(all_algorithm_files)
            del count_algs
        elif method_type == 's':
            all_algorithm_files = glob.glob(os.path.join(self.exp_output_path + '/history_result', '*.json'))
            all_algorithm_files = sorted(all_algorithm_files, key=lambda x: int(re.search(r'\d+', os.path.basename(x)).group()))
            for f in all_algorithm_files:
                with open(f) as file:
                    count_algs = json.load(file)['message']
                    self.algs_number += len(count_algs)
            del count_algs
        else:
            # warning here, if user give the wrong method type, the report may crash after long time run
            assert print("Method type not supported!")

    def create_convergence_pop(self):
        n_start = 0
        obj_list = np.zeros((self.pop_number, self.pop_size))
        all_algorithm_files = glob.glob(os.path.join(self.exp_output_path + '/history_result', '*.json'))
        all_algorithm_files = sorted(all_algorithm_files, key=lambda x: int(re.search(r'\d+', os.path.basename(x)).group()))

        for i, f in enumerate(all_algorithm_files):
            # Get result
            # Load JSON data from file
            with open(f) as file:
                data = json.load(file)['message']

            # Print each individual in the population
            na = 0
            for individual in data:
                code = individual['code']
                alg = individual['algorithm']
                obj = individual['objective']

                # code2file(alg,code,na,i)
                # print(obj)
                obj_list[i - n_start, na] = obj
                na += 1

        # Set font family to Times New Roman
        plt.rcParams['font.family'] = 'Times New Roman'
        # Generate x-axis values for number of generations
        generations = np.arange(1, obj_list.shape[0] + 1)
        best_objective = np.min(obj_list, axis=1)
        mean_objective = np.mean(obj_list, axis=1)

        # Set figure size
        plt.figure(figsize=(10, 6), dpi=80)

        # Plot objective value vs. number of generations for all samples as scatter points
        for i in generations:
            if i == 1:
                plt.scatter(i * np.ones(self.ini_pop_size), obj_list[i - 1, :self.ini_pop_size], color='tab:blue', alpha=0.6, s=200)
            else:
                plt.scatter(i * np.ones(self.pop_size), obj_list[i - 1, :], color='tab:blue', alpha=0.6, s=200)

        # Plot mean and best objectives
        plt.plot(generations, mean_objective, label='Mean', color='orange', linewidth=3.0)
        plt.plot(generations, best_objective, label='Best', color='r', linewidth=3.0)

        # Set plot title and labels with enlarged font size
        plt.xlabel('Number of Generations', fontsize=18)
        plt.ylabel('Obj.', fontsize=20)

        objmin = np.min(obj_list)
        objmax = np.max(obj_list)
        delta = (objmax - objmin) / 100.0
        # Set y-axis range
        plt.ylim([objmin - delta, objmax + delta])

        # Add scatter legend with enlarged font size
        plt.scatter([], [], color='tab:blue', alpha=0.6, label='Algorithms', s=200)  # Empty scatter plot for legend
        # plt.legend(scatterpoints=1, frameon=False, labelspacing=1, fontsize=20)
        plt.legend(scatterpoints=1, frameon=True, labelspacing=1, fontsize=20, fancybox=True, facecolor='gainsboro')
        # Adjust ticks and grid
        plt.xticks(np.arange(1, obj_list.shape[0] + 1, 2), fontsize=18)
        plt.yticks(np.arange(objmin - delta, objmax + delta, 10), fontsize=18)
        plt.grid(True, which='both', linestyle='--', linewidth=0.5)

        # Show the plot
        plt.tight_layout()
        plt.savefig(self.exp_output_path + '/report/ael_convergence.png')  # Save the plot as a file
        plt.savefig(self.exp_output_path + '/report/ael_convergence.pdf')
        # plt.show()

    def create_convergence_single(self):

        obj_list = np.zeros((self.algs_number, 2))
        all_algorithm_files = glob.glob(os.path.join(self.exp_output_path + '/history_result', '*.json'))
        all_algorithm_files = sorted(all_algorithm_files, key=lambda x: int(re.search(r'\d+', os.path.basename(x)).group()))

        na = 0
        for i, f in enumerate(all_algorithm_files):
            ### Get result ###
            # Load JSON data from file
            with open(f) as file:
                data = json.load(file)['message']

            # Print each individual in the population
            for individual in data:
                code = individual['code']
                alg = individual['algorithm']
                obj = individual['objective']

                # code2file(alg,code,na,i)
                # print(obj)
                obj_list[na, 0] = obj
                obj_list[na, 1] = np.mean(obj_list[:na + 1, 0])
                na += 1

        # Set font family to Times New Roman
        plt.rcParams['font.family'] = 'Times New Roman'
        # Generate x-axis values for number of generations
        generations = np.arange(1, obj_list.shape[0] + 1)
        best_objective = obj_list[:, 0]
        mean_objective = obj_list[:, 1]

        # Set figure size
        plt.figure(figsize=(10, 6), dpi=80)

        # Plot mean and best objectives
        plt.plot(generations, mean_objective, label='Mean', color='orange', linewidth=3.0)
        plt.plot(generations, best_objective, label='Best', color='r', linewidth=3.0)

        # Set plot title and labels with enlarged font size
        plt.xlabel('Number of Generations', fontsize=18)
        plt.ylabel('Obj.', fontsize=20)

        objmin = np.min(obj_list)
        objmax = np.max(obj_list)
        delta = (objmax - objmin) / 100.0
        # Set y-axis range
        plt.ylim([objmin - delta, objmax + delta])

        # plt.legend(scatterpoints=1, frameon=False, labelspacing=1, fontsize=20)
        plt.legend(scatterpoints=1, frameon=True, labelspacing=1, fontsize=20, fancybox=True, facecolor='gainsboro')
        # Adjust ticks and grid
        plt.xticks(np.arange(1, obj_list.shape[0] + 1, 2), fontsize=18)
        plt.yticks(np.arange(objmin - delta, objmax + delta, 10), fontsize=18)
        plt.grid(True, which='both', linestyle='--', linewidth=0.5)

        # Show the plot
        plt.tight_layout()
        plt.savefig(self.exp_output_path + '/report/ael_convergence.png')  # Save the plot as a file
        plt.savefig(self.exp_output_path + '/report/ael_convergence.pdf')
        # plt.show()

    def get_final_algorithms(self):
        ### Get result ###
        all_algorithm_files = glob.glob(os.path.join(self.exp_output_path + '/history_result', '*.json'))
        all_algorithm_files = sorted(all_algorithm_files, key=lambda x: int(re.search(r'\d+', os.path.basename(x)).group()))

        with open(all_algorithm_files[-1]) as file:
            data = json.load(file)['message']

        return data

    def generate_doc_report(self):
        # Create a new Document
        doc = Document()

        # Add Title
        doc.add_heading(self.title, level=1)

        # Add Parameter Settings
        doc.add_heading('Parameter Settings', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set Table Style
        table.style = 'Table Grid'

        # Set Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Value'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True  # Make headers bold

        # Add Data to Table
        for paras in self.paras:
            row_cells = table.add_row().cells
            row_cells[0].text = paras.__class__.__name__
            for attr, value in vars(paras).items():
                if attr != 'instance_data':
                    row_cells = table.add_row().cells
                    row_cells[0].text = attr
                    row_cells[1].text = str(value)
            if paras != self.paras[-1]:
                row = table.add_row().cells
                row[0].merge(row[1])

        # Add Borderlines
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)  # Set font size
                        run.font.bold = False  # Set font style
                        run.font.name = 'Calibri'  # Set font type
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color

        # Add Convergence Process
        doc.add_heading('Convergence Process', level=2)
        if self.method_type == 'p':
            self.create_convergence_pop()
        elif self.method_type == 's':
            self.create_convergence_single()
        else:
            assert print("Method type not supported!")
        doc.add_picture(self.exp_output_path + '/report/ael_convergence.png', width=Inches(4))

        # Add Final Results
        doc.add_heading('Final Results', level=2)
        algorithms_data = self.get_final_algorithms()

        # Add top algorithms data
        top_num = 1
        if self.method_type == 'p':
            top_num = 5
        elif self.method_type == 's':
            top_num = 1
        doc.add_heading(f'Top {top_num} Algorithms', level=3)
        for i, algorithm_data in enumerate(algorithms_data[:top_num]):
            doc.add_heading(f'Algorithm {i + 1}', level=4)
            doc.add_paragraph(f'Algorithm: {algorithm_data["algorithm"]}')
            # Create a new paragraph
            p = doc.add_paragraph()

            # Add the code block with background color and border
            code = algorithm_data["code"]
            code_block = p.add_run()
            code_block.text = f'Code:\n{code}'
            code_block_font = code_block.font
            code_block_font.size = Pt(8)

            # Set the background color
            shading_elm = OxmlElement('w:shd')
            shading_elm.set('fill', 'D9D9D9')  # Set your desired background color here
            p._element.append(shading_elm)

            # Set the border
            # p.border_top.space = Pt(1)     # Set border space

            doc.add_paragraph(f'Fitness: {algorithm_data["objective"]}')
            doc.add_paragraph('')  # Add a blank paragraph for separation

        # Save the document
        doc.save(self.exp_output_path + '/report/ael_report.docx')


class LogData:

    def __init__(self, paras, name, log_type='txt', method_type='p', detailed=False):
        os.environ['OMP_NUM_THREADS'] = '1'
        self.__method_type = method_type
        self.__log_type = log_type
        self.__paras = paras  # method, problem, etc.
        self.__process_start_time = process_start_time
        self.__json_path = None
        self.__report_path = None

        self.__result_folder = result_folder
        self.__result_folder_ini = result_folder_ini  # backup

        # specific log storage parameters
        log_name = 'runlog.txt'
        if self.__log_type == 'txt':
            log_name = 'runlog.txt'
            self.logger = logging.getLogger('txt')
        elif self.__log_type == 'json':
            log_name = 'runlog.json'
            self.logger = logging.getLogger('json')
        else:
            assert print("Log type not supported!")

        self.__used_txt_file = [log_name]
        self.__used_json_file = []
        self.is_detailed = detailed
        self.logger_params = {
            'log_file': {
                'filepath': self.get_result_folder(),
                'desc': name,
                'filename': log_name
            }
        }

        # result record
        self.history = []

        # initialize
        self.initialize()

    def initialize(self):
        self._create_logger(**self.logger_params)

        if self.__log_type == 'txt':
            self._copy_all_src(self.get_result_folder())
            self.print_config(self.logger)
            self.set_write_file_a(file_n='debug_log.txt', level=logging.DEBUG)
        if self.__log_type == 'json':
            self.set_write_file_r("runlog.json")
            self.set_write_json_a(file_n='debug_log.json', result_type='debug', level=logging.DEBUG)
        self.set_result_folder_ini()

        if self.is_detailed:
            self._set_write_terminal(level=logging.DEBUG)

    # print initial configuration
    def print_config(self, logger):
        logger.info("==================================Parameters===============================")
        for p in self.__paras:
            if p != self.__paras[0]:
                logger.info("=================================================================")
            logger.info(f"{p.__class__.__name__}: ")
            for attr, value in p.__dict__.items():
                if attr not in ['instance_data', 'llm', 'prob', 'txt_logger', 'json_logger']:
                    logger.info(f"{attr}: {value}")
        [logger.info(g_key + "{}".format(self.logger_params[g_key])) for g_key in self.logger_params.keys()]
        logger.info("==================================Parameters===============================")

    # tools for log data
    def get_result_folder(self):
        return self.__result_folder

    def set_result_folder(self, folder):
        self.__result_folder = folder

    def set_result_folder_ini(self):
        self.__result_folder = self.__result_folder_ini

    def return_set_files(self):
        file_handler = []
        for handler in self.logger.handlers:
            if type(handler) is logging.FileHandler:
                file_handler.append(handler)

        return file_handler

    # when set to logging.debug, debug info will be printed
    def _set_write_terminal(self, to_terminal=True, level=logging.INFO):

        console_handler = None
        for handler in self.logger.handlers:
            if type(handler) is logging.StreamHandler:
                console_handler = handler
                break

        if to_terminal:
            if not console_handler:
                formatter = logging.Formatter("[%(asctime)s] %(filename)s(%(lineno)d) : %(message)s",
                                              "%Y-%m-%d %H:%M:%S")
                console = logging.StreamHandler(sys.stdout)
                console.setLevel(level)
                console.setFormatter(formatter)
                self.logger.addHandler(console)
        else:
            if console_handler:
                console_handler.flush()
                self.logger.removeHandler(console_handler)

    def only_write_terminal(self, level=logging.INFO):

        formatter = logging.Formatter("[%(asctime)s] %(filename)s(%(lineno)d) : %(message)s",
                                      "%Y-%m-%d %H:%M:%S")
        console = logging.StreamHandler(sys.stdout)
        console.setLevel(level)
        console.setFormatter(formatter)

        # remove other file handler
        remove_handler = []
        for handler in self.logger.handlers:
            if type(handler) is not logging.StreamHandler:
                remove_handler.append(handler)
        for handler in remove_handler:
            self.logger.removeHandler(handler)

        self.logger.addHandler(console)

    def set_write_file_a(self, file_n="default_log.txt", level=logging.INFO):
        if file_n not in self.__used_txt_file:
            self.__used_txt_file.append(file_n)

        logger_file = self.logger_params
        logger_file['log_file']['filename'] = file_n
        logger_file['log_file']['filepath'] = logger_file['log_file']['filepath'].format(desc='_' + logger_file['log_file']['desc'])
        filename = logger_file['log_file']['filepath'] + '/' + logger_file['log_file']['filename']

        if not os.path.exists(logger_file['log_file']['filepath']):
            os.makedirs(logger_file['log_file']['filepath'])

        file_mode = 'a' if os.path.isfile(filename) else 'w'
        formatter = logging.Formatter("[%(asctime)s] %(filename)s(%(lineno)d) : %(message)s",
                                      "%Y-%m-%d %H:%M:%S")
        fileout = logging.FileHandler(filename, mode=file_mode)
        fileout.setLevel(level)
        fileout.setFormatter(formatter)
        self.logger.addHandler(fileout)

    def set_write_file_r(self, file_n):
        file_handler = []
        for handler in self.logger.handlers:
            if type(handler) is logging.FileHandler and os.path.basename(handler.baseFilename) == file_n:
                file_handler.append(handler)

        if file_handler:
            for fh in file_handler:
                # 如果已有文件处理器但不需要输出到文件，则移除文件处理器
                fh.flush()
                self.logger.removeHandler(fh)
        else:
            assert print("File do not exist！")

    def denote_write_file(self, file_n, level=logging.INFO):
        if file_n not in self.__used_txt_file:
            self.__used_txt_file.append(file_n)

        logger_file = self.logger_params
        logger_file['log_file']['filename'] = file_n
        logger_file['log_file']['filepath'] = logger_file['log_file']['filepath'].format(
            desc='_' + logger_file['log_file']['desc'])
        filename = logger_file['log_file']['filepath'] + '/' + logger_file['log_file']['filename']

        if not os.path.exists(logger_file['log_file']['filepath']):
            os.makedirs(logger_file['log_file']['filepath'])

        file_mode = 'a' if os.path.isfile(filename) else 'w'
        formatter = logging.Formatter("[%(asctime)s] %(filename)s(%(lineno)d) : %(message)s",
                                      "%Y-%m-%d %H:%M:%S")
        fileout = logging.FileHandler(filename, mode=file_mode)
        fileout.setLevel(level)
        fileout.setFormatter(formatter)

        # remove other file handler
        remove_handler = []
        for handler in self.logger.handlers:
            if type(handler) is logging.FileHandler:  # and os.path.basename(handler.baseFilename) != file_n:
                remove_handler.append(handler)
        for handler in remove_handler:
            self.logger.removeHandler(handler)

        self.logger.addHandler(fileout)

    def delete_defined_file(self, file_n):
        file_handler = []
        for handler in self.logger.handlers:
            if type(handler) is logging.FileHandler and os.path.basename(handler.baseFilename) == file_n:
                file_handler.append(handler)

        if file_handler:
            for fh in file_handler:
                # 如果已有文件处理器但不需要输出到文件，则移除文件处理器
                self.logger.removeHandler(fh)
                fh.flush()
                fh.close()
        else:
            assert print("File do not exist！")

    def use_all_defined_file(self):
        for f in self.__used_txt_file:
            self.set_write_file_a(f)

    def denote_write_json(self, file_n, result_type='history', level=logging.INFO):
        if file_n not in self.__used_json_file:
            self.__used_json_file.append(file_n)

        logger_file = self.logger_params
        logger_file['log_file']['filename'] = file_n
        logger_file['log_file']['filepath'] = logger_file['log_file']['filepath'].format(desc='_' + logger_file['log_file']['desc'])

        if result_type == 'history':
            filename = logger_file['log_file']['filepath'] + '/' + 'history_result/' + logger_file['log_file']['filename']
        elif result_type == 'best':
            filename = logger_file['log_file']['filepath'] + '/' + 'best_result/' + logger_file['log_file']['filename']
        elif result_type == 'debug':
            filename = logger_file['log_file']['filepath'] + '/' + 'debug_result/' + logger_file['log_file']['filename']
        else:
            assert print("Type not supported!")

        # if not os.path.exists(logger_file['log_file']['filepath']):
        #     os.makedirs(logger_file['log_file']['filepath'])

        file_mode = 'a' if os.path.isfile(filename) else 'w'
        formatter = JSONFormatter()
        fileout = logging.FileHandler(filename, mode=file_mode)
        fileout.setLevel(level)
        fileout.setFormatter(formatter)

        # remove other file handler
        remove_handler = []
        for handler in self.logger.handlers:
            if type(handler) is logging.FileHandler:  #  and os.path.basename(handler.baseFilename) != file_n:
                remove_handler.append(handler)
        for handler in remove_handler:
            self.logger.removeHandler(handler)

        self.logger.addHandler(fileout)

    def set_write_json_a(self, file_n="default_log.json", result_type='history', level=logging.INFO):
        if file_n not in self.__used_json_file:
            self.__used_json_file.append(file_n)

        logger_file = self.logger_params
        logger_file['log_file']['filename'] = file_n
        logger_file['log_file']['filepath'] = logger_file['log_file']['filepath'].format(desc='_' + logger_file['log_file']['desc'])

        if result_type == 'history':
            filename = logger_file['log_file']['filepath'] + '/' + 'history_result/' + logger_file['log_file']['filename']
        elif result_type == 'best':
            filename = logger_file['log_file']['filepath'] + '/' + 'best_result/' + logger_file['log_file']['filename']
        elif result_type == 'debug':
            filename = logger_file['log_file']['filepath'] + '/' + 'debug_result/' + logger_file['log_file']['filename']
        else:
            assert print("Type not supported!")

        # if not os.path.exists(logger_file['log_file']['filepath']):
        #     os.makedirs(logger_file['log_file']['filepath'])

        file_mode = 'a' if os.path.isfile(filename) else 'w'
        formatter = JSONFormatter()
        fileout = logging.FileHandler(filename, mode=file_mode)
        fileout.setLevel(level)
        fileout.setFormatter(formatter)
        self.logger.addHandler(fileout)

    def delete_defined_json(self, file_n):
        file_handler = []
        for handler in self.logger.handlers:
            if type(handler) is logging.FileHandler and os.path.basename(handler.baseFilename) == file_n:
                file_handler.append(handler)

        if file_handler:
            for fh in file_handler:
                # 如果已有文件处理器但不需要输出到文件，则移除文件处理器
                self.logger.removeHandler(fh)
                fh.flush()
                fh.close()
        else:
            assert print("File do not exist！")

    def set_write_json_r(self, file_n):
        file_handler = []
        for handler in self.logger.handlers:
            if type(handler) is logging.FileHandler and os.path.basename(handler.baseFilename) == file_n:
                file_handler.append(handler)

        if file_handler:
            for fh in file_handler:
                # 如果已有文件处理器但不需要输出到文件，则移除文件处理器
                fh.flush()
                self.logger.removeHandler(fh)
        else:
            assert print("File do not exist！")

    def use_all_defined_json(self):
        for f in self.__used_json_file:
            self.set_write_json_a(f)

    def _create_logger(self, log_file=None):
        if 'filepath' not in log_file:
            log_file['filepath'] = self.get_result_folder()

        if 'desc' in log_file:
            log_file['filepath'] = log_file['filepath'].format(desc='_' + log_file['desc'])
        else:
            log_file['filepath'] = log_file['filepath'].format(desc='')

        self.set_result_folder(log_file['filepath'])

        if 'filename' in log_file:
            filename = log_file['filepath'] + '/' + log_file['filename']
        else:
            filename = log_file['filepath'] + '/' + 'log.txt'

        # create folder
        if not os.path.exists(log_file['filepath']):
            os.makedirs(log_file['filepath'])
            self._create_result_folder(log_file['filepath'])
            self._create_report_folder(log_file['filepath'])

        file_mode = 'a' if os.path.isfile(filename) else 'w'

        # self.logger = logging.getLogger()
        self.logger.setLevel(level=logging.INFO)

        formatter = logging.Formatter("[%(asctime)s] %(filename)s(%(lineno)d) : %(message)s", "%Y-%m-%d %H:%M:%S")
        if self.__log_type == 'json':
            formatter = JSONFormatter()
        elif self.__log_type == 'txt':
            formatter = logging.Formatter("[%(asctime)s] %(filename)s(%(lineno)d) : %(message)s", "%Y-%m-%d %H:%M:%S")
        else:
            assert print("Log type not supported!")

        for hdlr in self.logger.handlers[:]:
            self.logger.removeHandler(hdlr)

        # write to file
        fileout = logging.FileHandler(filename, mode=file_mode)
        fileout.setLevel(logging.INFO)
        fileout.setFormatter(formatter)
        self.logger.addHandler(fileout)

        # write to console
        if self.__log_type == 'txt':
            console = logging.StreamHandler(sys.stdout)
            console.setLevel(logging.INFO)
            console.setFormatter(formatter)
            self.logger.addHandler(console)

    def _create_result_folder(self, dst_root):

        # make target directory
        best_result_path = os.path.join(dst_root, 'best_result')
        history_result_path = os.path.join(dst_root, 'history_result')
        debug_result_path = os.path.join(dst_root, 'debug_result')

        os.makedirs(best_result_path)
        os.makedirs(history_result_path)
        os.makedirs(debug_result_path)

    def _create_report_folder(self, dst_root):
        # make target directory
        report_path = os.path.join(dst_root, 'report')

        os.makedirs(report_path)

    def _copy_all_src(self, dst_root):
        # execution dir
        if os.path.basename(sys.argv[0]).startswith('ipykernel_launcher'):
            execution_path = os.getcwd()
        else:
            execution_path = os.path.dirname(sys.argv[0])

        # home dir setting
        tmp_dir1 = os.path.abspath(os.path.join(execution_path, sys.path[-2]))
        tmp_dir2 = os.path.abspath(os.path.join(execution_path, sys.path[-1]))

        if len(tmp_dir1) > len(tmp_dir2) and os.path.exists(tmp_dir2):
            home_dir = tmp_dir2
        else:
            home_dir = tmp_dir1

        # make target directory
        dst_path = os.path.join(dst_root, 'src')

        if not os.path.exists(dst_path):
            os.makedirs(dst_path)

        for item in sys.modules.items():
            key, value = item

            if hasattr(value, '__file__') and value.__file__:
                src_abspath = os.path.abspath(value.__file__)

                if os.path.commonprefix([home_dir, src_abspath]) == home_dir:
                    dst_filepath = os.path.join(dst_path, os.path.basename(src_abspath))

                    if os.path.exists(dst_filepath):
                        split = list(os.path.splitext(dst_filepath))
                        split.insert(1, '({})')
                        filepath = ''.join(split)
                        post_index = 0

                        while os.path.exists(filepath.format(post_index, desc='{desc}')):
                            post_index += 1

                        dst_filepath = filepath.format(post_index, desc='{desc}')

                    shutil.copy(src_abspath, dst_filepath)

    def gen_analyse_report(self):
        history_mean = np.mean(self.history)
        history_std = np.std(self.history)
        title = self.__paras[2].__class__.__name__ + "_" + self.__paras[1].__class__.__name__ + "_" + process_start_time.strftime("%Y%m%d_%H%M%S")

        report = ReportCreator(self.__paras, title, self.logger_params['log_file']['filepath'], method_type=self.__method_type)
        report.generate_doc_report()
